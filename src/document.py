"""Document word document class"""

import time
import docx
import os


class Document:
    """Document class"""

    def __init__(self):
        self.doc = docx.Document()

    def save_document(self, name: str = ""):
        """save document to output"""

        if name != "":
            os.path.join(os.getcwd(), f'/out/{name}.docx')
            self.doc.save(os.path.join(os.getcwd(), f'out/{name}.docx'))
        else:
            self.doc.save(os.path.join(
                os.getcwd(), f'out/out_{time.strftime("%Y%m%d_%H%M%S")}.docx'))

    def add_heading(self, header: str = "", level: int = 1):
        """add heading to document

        Args:
            header (str, optional): Defaults to "".
            level (int, optional): Defaults to 1.

        Returns:
            _type_: _description_
        """

        return self.doc.add_heading(header, level=level)

    def add_paragraphe(self, line: str = ""):
        """add paragraphe to document

        Args:
            line (str, optional): Defaults to "".

        Returns:
            _type_: _description_
        """

        return self.doc.add_paragraph(line)

    def add_ordered_list(self, line: str = ""):
        """add ordered list to document

        Args:
            line (str, optional): Defaults to "".

        Returns:
            _type_: _description_
        """

        return self.doc.add_paragraph(line, style="List Number")

    def add_unordered_list(self, line: str = ""):
        """add ordered list to document

        Args:
            line (str, optional): Defaults to "".

        Returns:
            _type_: _description_
        """

        return self.doc.add_paragraph(line, style="List Bullet")

    def add_table(self, data: list, rows: int = 1, cols: int = 1):
        """add table to document

        Args:
            rows (int, optional): Defaults to 1.
            cols (int, optional): Defaults to 1.
            data (list, optional): Defaults to [[]].

        Returns:
            _type_: _description_
        """

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        if data and len(data):
            for row in data:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = cell
        return table
